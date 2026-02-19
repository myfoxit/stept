"""
Utility functions for exporting documents/pages to various formats.
Handles TipTap JSON content conversion without brittle manual pagination.
"""
import io
import os
import httpx
import html as html_module
from typing import Any, Dict, List, Optional

# Gotenberg configuration
GOTENBERG_URL = os.getenv("GOTENBERG_URL", "http://gotenberg:3000")

# Page dimensions in mm (matching frontend PAGE_FORMATS)
# Frontend uses pixels at ~96dpi, so we convert:
# A4: 794px width ≈ 210mm, margins typically 96px ≈ 25.4mm (1 inch)
# Letter: 816px width ≈ 215.9mm
# Note: Frontend PAGE_FORMATS uses { width, height, margins: { top, bottom, left, right } } in pixels
# At 96dpi: 1 inch = 96px = 25.4mm
# Typical frontend margins: 96px = 1 inch = 25.4mm
PAGE_FORMATS = {
    "a4": {
        "width_mm": 210,
        "height_mm": 297,
        # Derived from browser pagination/index.ts PAGE_FORMATS.A4:
        # width=794px, height=1123px (cmToPixels(21, 96), cmToPixels(29.7, 96))
        # margins: left/right=76px (2cm), top/bottom=94px (2.5cm)
        # Plus 1px border on each side (border-box) → effective margin = padding + 1
        # Paper: 794/96 in, margins: 77/96 in (left/right), 95/96 in (top/bottom)
        # Content: 794 - 77*2 = 640px wide, 1123 - 95*2 = 933px tall
        "width_in": 794 / 96,
        "height_in": 1123 / 96,
        "margins_in": {
            "top": 95 / 96,      # 94px padding + 1px border
            "bottom": 95 / 96,
            "left": 77 / 96,     # 76px padding + 1px border
            "right": 77 / 96,
        },
    },
    "a3": {
        "width_mm": 297,
        "height_mm": 420,
        # Browser: 1123x1587px, margins t/b=94px l/r=76px, +1px border → content 969x1397px
        "width_in": 1123 / 96,
        "height_in": 1587 / 96,
        "margins_in": {
            "top": 95 / 96,
            "bottom": 95 / 96,
            "left": 77 / 96,
            "right": 77 / 96,
        },
    },
    "a5": {
        "width_mm": 148,
        "height_mm": 210,
        # Browser: 559x794px, margins t/b=76px l/r=57px, +1px border → content 443x640px
        "width_in": 559 / 96,
        "height_in": 794 / 96,
        "margins_in": {
            "top": 77 / 96,
            "bottom": 77 / 96,
            "left": 58 / 96,
            "right": 58 / 96,
        },
    },
    "letter": {
        "width_mm": 215.9,
        "height_mm": 279.4,
        # Browser: 816x1056px, margins t/b/l/r=96px, +1px border → content 622x862px
        "width_in": 816 / 96,
        "height_in": 1056 / 96,
        "margins_in": {
            "top": 97 / 96,
            "bottom": 97 / 96,
            "left": 97 / 96,
            "right": 97 / 96,
        },
    },
    "legal": {
        "width_mm": 215.9,
        "height_mm": 355.6,
        # Browser: 816x1344px, margins t/b/l/r=96px, +1px border → content 622x1150px
        "width_in": 816 / 96,
        "height_in": 1344 / 96,
        "margins_in": {
            "top": 97 / 96,
            "bottom": 97 / 96,
            "left": 97 / 96,
            "right": 97 / 96,
        },
    },
    "tabloid": {
        "width_mm": 279.4,
        "height_mm": 431.8,
        # Browser: 1056x1632px, margins t/b/l/r=96px, +1px border → content 862x1438px
        "width_in": 1056 / 96,
        "height_in": 1632 / 96,
        "margins_in": {
            "top": 97 / 96,
            "bottom": 97 / 96,
            "left": 97 / 96,
            "right": 97 / 96,
        },
    },
}


def tiptap_to_markdown(content: Any) -> str:
    """Convert TipTap JSON content to Markdown."""
    if not content:
        return ""
    
    if isinstance(content, str):
        return content
    
    def process_marks(text: str, marks: List[Dict]) -> str:
        for mark in marks:
            mark_type = mark.get("type", "")
            if mark_type == "bold":
                text = f"**{text}**"
            elif mark_type == "italic":
                text = f"*{text}*"
            elif mark_type == "strike":
                text = f"~~{text}~~"
            elif mark_type == "code":
                text = f"`{text}`"
            elif mark_type == "link":
                href = mark.get("attrs", {}).get("href", "")
                text = f"[{text}]({href})"
        return text
    
    def convert_node(node: Dict[str, Any], list_depth: int = 0, ordered_counter: Optional[int] = None) -> str:
        node_type = node.get("type", "")
        result = ""
        
        if node_type == "doc":
            for child in node.get("content", []):
                result += convert_node(child)
            return result
        
        if node_type == "text":
            text = node.get("text", "")
            marks = node.get("marks", [])
            return process_marks(text, marks)
        
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            prefix = "#" * level + " "
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"\n{prefix}{content}\n\n"
        
        if node_type == "paragraph":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"{content}\n\n"
        
        if node_type == "bulletList":
            for child in node.get("content", []):
                result += convert_node(child, list_depth + 1)
            return result
        
        if node_type == "orderedList":
            counter = 1
            for child in node.get("content", []):
                result += convert_node(child, list_depth + 1, counter)
                counter += 1
            return result
        
        if node_type == "listItem":
            indent = "  " * (list_depth - 1)
            if ordered_counter:
                prefix = f"{indent}{ordered_counter}. "
            else:
                prefix = f"{indent}- "
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"{prefix}{content.strip()}\n"
        
        if node_type == "codeBlock":
            lang = node.get("attrs", {}).get("language", "")
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"\n```{lang}\n{content}\n```\n\n"
        
        if node_type == "blockquote":
            content = "".join(convert_node(c) for c in node.get("content", []))
            lines = content.strip().split("\n")
            quoted = "\n".join(f"> {line}" for line in lines)
            return f"\n{quoted}\n\n"
        
        if node_type == "horizontalRule":
            return "\n---\n\n"
        
        if node_type == "hardBreak":
            return "  \n"
        
        if node_type == "image":
            src = node.get("attrs", {}).get("src", "")
            alt = node.get("attrs", {}).get("alt", "")
            return f"![{alt}]({src})\n\n"
        
        if node_type == "pageBreak":
            return "\n\n---\n*[Page Break]*\n---\n\n"
        
        if node_type == "taskList":
            for child in node.get("content", []):
                result += convert_node(child, list_depth + 1)
            return result
        
        if node_type == "taskItem":
            indent = "  " * (list_depth - 1)
            checked = node.get("attrs", {}).get("checked", False)
            checkbox = "[x]" if checked else "[ ]"
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"{indent}- {checkbox} {content.strip()}\n"
        
        if node_type == "table":
            rows = node.get("content", [])
            if not rows:
                return ""
            table_md = []
            for i, row in enumerate(rows):
                cells = row.get("content", [])
                cell_texts = []
                for cell in cells:
                    cell_content = "".join(convert_node(c) for c in cell.get("content", []))
                    cell_texts.append(cell_content.strip().replace("\n", " "))
                table_md.append("| " + " | ".join(cell_texts) + " |")
                if i == 0:
                    table_md.append("| " + " | ".join(["---"] * len(cell_texts)) + " |")
            return "\n" + "\n".join(table_md) + "\n\n"
        
        if node_type == "process-recording-node":
            session_id = node.get("attrs", {}).get("sessionId", "unknown")
            return f"\n[Embedded Workflow Recording: {session_id}]\n\n"
        
        if node_type == "dataTable":
            table_name = node.get("attrs", {}).get("tableName", "")
            table_id = node.get("attrs", {}).get("tableId", "unknown")
            label = table_name or table_id
            return f"\n[Embedded Data Table: {label}]\n\n"
        
        for child in node.get("content", []):
            result += convert_node(child, list_depth, ordered_counter)
        return result
    
    if isinstance(content, dict):
        return convert_node(content).strip()
    
    return ""


def tiptap_nodes_to_html(nodes: List[Dict], page_layout: str = "document") -> str:
    """Convert a list of TipTap nodes to HTML."""
    
    def process_marks(text: str, marks: List[Dict]) -> str:
        for mark in marks:
            mark_type = mark.get("type", "")
            if mark_type == "bold":
                text = f"<strong>{text}</strong>"
            elif mark_type == "italic":
                text = f"<em>{text}</em>"
            elif mark_type == "strike":
                text = f"<s>{text}</s>"
            elif mark_type == "code":
                text = f"<code>{text}</code>"
            elif mark_type == "link":
                href = mark.get("attrs", {}).get("href", "")
                target = mark.get("attrs", {}).get("target", "_blank")
                text = f'<a href="{href}" target="{target}">{text}</a>'
            elif mark_type == "underline":
                text = f"<u>{text}</u>"
        return text
    
    def convert_node(node: Dict[str, Any]) -> str:
        node_type = node.get("type", "")
        
        if node_type == "doc":
            return "".join(convert_node(c) for c in node.get("content", []))
        
        if node_type == "text":
            text = html_module.escape(node.get("text", ""))
            marks = node.get("marks", [])
            return process_marks(text, marks)
        
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<h{level}>{content}</h{level}>"
        
        if node_type == "paragraph":
            content = "".join(convert_node(c) for c in node.get("content", []))
            # Preserve textAlign attribute
            text_align = node.get("attrs", {}).get("textAlign", "") if node.get("attrs") else ""
            style = f' style="text-align: {text_align}"' if text_align and text_align != "left" else ""
            # Preserve empty paragraphs for vertical spacing
            if not content:
                return f"<p{style}><br></p>"
            return f"<p{style}>{content}</p>"
        
        if node_type == "bulletList":
            items = "".join(convert_node(c) for c in node.get("content", []))
            return f"<ul>{items}</ul>"
        
        if node_type == "orderedList":
            items = "".join(convert_node(c) for c in node.get("content", []))
            start = node.get("attrs", {}).get("start", 1)
            return f'<ol start="{start}">{items}</ol>'
        
        if node_type == "listItem":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<li>{content}</li>"
        
        if node_type == "codeBlock":
            lang = node.get("attrs", {}).get("language", "")
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f'<pre><code class="language-{lang}">{content}</code></pre>'
        
        if node_type == "blockquote":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<blockquote>{content}</blockquote>"
        
        if node_type == "horizontalRule":
            return "<hr>"
        
        # Explicit Page Breaks
        if node_type == "pageBreak":
            return '<div style="break-after: page; page-break-after: always;"></div>'

        if node_type == "hardBreak":
            return "<br>"
        
        if node_type == "image":
            src = node.get("attrs", {}).get("src", "")
            alt = node.get("attrs", {}).get("alt", "")
            return f'<img src="{src}" alt="{alt}" style="max-width: 100%; display: block; margin: 10px 0;">'
        
        if node_type == "table":
            rows = "".join(convert_node(c) for c in node.get("content", []))
            return f"<table>{rows}</table>"
        
        if node_type == "tableRow":
            cells = "".join(convert_node(c) for c in node.get("content", []))
            return f"<tr>{cells}</tr>"
        
        if node_type == "tableCell":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<td>{content}</td>"
        
        if node_type == "tableHeader":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<th>{content}</th>"
        
        if node_type == "taskList":
            items = "".join(convert_node(c) for c in node.get("content", []))
            return f'<ul style="list-style: none; padding-left: 0;">{items}</ul>'
        
        if node_type == "taskItem":
            checked = node.get("attrs", {}).get("checked", False)
            checkbox = "☑" if checked else "☐"
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<li>{checkbox} {content}</li>"
        
        if node_type == "process-recording-node":
            session_id = node.get("attrs", {}).get("sessionId", "unknown")
            return f'<div class="embedded-node">[Embedded Workflow Recording: {html_module.escape(str(session_id))}]</div>'
        
        if node_type == "dataTable":
            table_name = node.get("attrs", {}).get("tableName", "")
            table_id = node.get("attrs", {}).get("tableId", "unknown")
            label = table_name or table_id
            return f'<div class="embedded-node">[Embedded Data Table: {html_module.escape(str(label))}]</div>'
        
        return "".join(convert_node(c) for c in node.get("content", []))
    
    return "".join(convert_node(node) for node in nodes)


def tiptap_to_html(content: Any, page_layout: str = "document") -> str:
    """Convert TipTap JSON content to HTML body content."""
    if not content:
        return ""
    
    if isinstance(content, str):
        return f"<p>{content}</p>"
    
    if isinstance(content, dict):
        nodes = content.get("content", [])
        return tiptap_nodes_to_html(nodes, page_layout)
    
    return ""


def generate_document_confluence(doc: Any, page_layout: str = "document") -> str:
    """Generate Confluence Storage Format export of a document."""
    if not doc.content:
        return ""

    def convert_node(node: Dict[str, Any]) -> str:
        node_type = node.get("type", "")

        if node_type == "doc":
            return "".join(convert_node(c) for c in node.get("content", []))

        if node_type == "text":
            text = html_module.escape(node.get("text", ""))
            for mark in node.get("marks", []):
                mt = mark.get("type", "")
                if mt == "bold":
                    text = f"<strong>{text}</strong>"
                elif mt == "italic":
                    text = f"<em>{text}</em>"
                elif mt == "code":
                    text = f"<code>{text}</code>"
                elif mt == "link":
                    href = mark.get("attrs", {}).get("href", "")
                    text = f'<a href="{href}">{text}</a>'
            return text

        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<h{level}>{content}</h{level}>"

        if node_type == "paragraph":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<p>{content}</p>" if content else "<p></p>"

        if node_type == "bulletList":
            items = "".join(convert_node(c) for c in node.get("content", []))
            return f"<ul>{items}</ul>"

        if node_type == "orderedList":
            items = "".join(convert_node(c) for c in node.get("content", []))
            return f"<ol>{items}</ol>"

        if node_type == "listItem":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<li>{content}</li>"

        if node_type == "codeBlock":
            lang = node.get("attrs", {}).get("language", "")
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f'<ac:structured-macro ac:name="code"><ac:parameter ac:name="language">{lang}</ac:parameter><ac:plain-text-body><![CDATA[{content}]]></ac:plain-text-body></ac:structured-macro>'

        if node_type == "blockquote":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<blockquote>{content}</blockquote>"

        if node_type == "horizontalRule":
            return "<hr/>"

        if node_type == "image":
            src = node.get("attrs", {}).get("src", "")
            return f'<ac:image><ri:url ri:value="{src}"/></ac:image>'

        if node_type == "table":
            rows = "".join(convert_node(c) for c in node.get("content", []))
            return f"<table>{rows}</table>"

        if node_type in ("tableRow", "tableCell", "tableHeader"):
            tag = {"tableRow": "tr", "tableCell": "td", "tableHeader": "th"}[node_type]
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"<{tag}>{content}</{tag}>"

        if node_type == "process-recording-node":
            session_id = node.get("attrs", {}).get("sessionId", "unknown")
            return f"<p><em>[Embedded Workflow Recording: {html_module.escape(str(session_id))}]</em></p>"

        if node_type == "dataTable":
            table_name = node.get("attrs", {}).get("tableName", "")
            table_id = node.get("attrs", {}).get("tableId", "unknown")
            label = table_name or table_id
            return f"<p><em>[Embedded Data Table: {html_module.escape(str(label))}]</em></p>"

        return "".join(convert_node(c) for c in node.get("content", []))

    if isinstance(doc.content, dict):
        return convert_node(doc.content)
    return ""


def generate_document_notion_markdown(doc: Any, page_layout: str = "document") -> str:
    """Generate Notion-compatible Markdown export of a document."""
    if not doc.content:
        return ""

    def convert_node(node: Dict[str, Any], list_depth: int = 0, ordered_counter: Optional[int] = None) -> str:
        node_type = node.get("type", "")

        if node_type == "doc":
            return "".join(convert_node(c) for c in node.get("content", []))

        if node_type == "text":
            text = node.get("text", "")
            for mark in node.get("marks", []):
                mt = mark.get("type", "")
                if mt == "bold":
                    text = f"**{text}**"
                elif mt == "italic":
                    text = f"*{text}*"
                elif mt == "code":
                    text = f"`{text}`"
                elif mt == "link":
                    href = mark.get("attrs", {}).get("href", "")
                    text = f"[{text}]({href})"
            return text

        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"\n{'#' * level} {content}\n\n"

        if node_type == "paragraph":
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"{content}\n\n"

        if node_type == "bulletList":
            result = ""
            for child in node.get("content", []):
                result += convert_node(child, list_depth + 1)
            return result

        if node_type == "orderedList":
            counter = 1
            result = ""
            for child in node.get("content", []):
                result += convert_node(child, list_depth + 1, counter)
                counter += 1
            return result

        if node_type == "listItem":
            indent = "  " * (list_depth - 1)
            prefix = f"{indent}{ordered_counter}. " if ordered_counter else f"{indent}- "
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"{prefix}{content.strip()}\n"

        if node_type == "codeBlock":
            lang = node.get("attrs", {}).get("language", "")
            content = "".join(convert_node(c) for c in node.get("content", []))
            return f"\n```{lang}\n{content}\n```\n\n"

        if node_type == "blockquote":
            content = "".join(convert_node(c) for c in node.get("content", []))
            lines = content.strip().split("\n")
            quoted = "\n".join(f"> {line}" for line in lines)
            return f"\n> 💡 {quoted.lstrip('> ')}\n\n"

        if node_type == "horizontalRule":
            return "\n---\n\n"

        if node_type == "image":
            src = node.get("attrs", {}).get("src", "")
            alt = node.get("attrs", {}).get("alt", "")
            return f"![{alt}]({src})\n\n"

        if node_type == "process-recording-node":
            session_id = node.get("attrs", {}).get("sessionId", "unknown")
            return f"\n[Embedded Workflow Recording: {session_id}]\n\n"

        if node_type == "dataTable":
            table_name = node.get("attrs", {}).get("tableName", "")
            table_id = node.get("attrs", {}).get("tableId", "unknown")
            label = table_name or table_id
            return f"\n[Embedded Data Table: {label}]\n\n"

        result = ""
        for child in node.get("content", []):
            result += convert_node(child, list_depth, ordered_counter)
        return result

    if isinstance(doc.content, dict):
        return convert_node(doc.content).strip()
    return ""


def generate_document_markdown(doc: Any, page_layout: str = "document") -> str:
    """Generate Markdown export of a document."""
    lines = []
    if doc.content:
        md_content = tiptap_to_markdown(doc.content)
        lines.append(md_content)
    return "\n".join(lines)


def generate_document_html(
    doc: Any, 
    embed_styles: bool = True, 
    page_layout: str = "document",
    for_pdf: bool = False
) -> str:
    """Generate HTML export of a document."""
    title = doc.name or "Untitled Document"
    safe_title = html_module.escape(title)
    
    # PDF Styles: Gotenberg handles page sizes and margins via its API parameters.
    # We only add print-specific styles to avoid awkward breaks.
    # IMPORTANT: Do NOT set @page { margin: 0 } as it overrides Gotenberg's margins!
    pdf_styles = """
        body {
            margin: 0;
            padding: 0;
            word-wrap: break-word;
        }
        /* Prevent awkward breaks */
        p, blockquote, pre, table, figure, li {
            page-break-inside: avoid;
            break-inside: avoid;
        }
        h1, h2, h3, h4, h5, h6 {
            page-break-after: avoid;
            break-after: avoid;
        }
        table { width: 100% !important; }
    """ if for_pdf else ""
    
    base_styles = """
        *, *::before, *::after { box-sizing: border-box; }
        body { 
            font-family: Arial, sans-serif; 
            font-size: 1rem;
            line-height: 1.6;
            color: #1e293b;
            text-rendering: optimizeLegibility;
            -webkit-font-smoothing: antialiased;
            overflow-wrap: break-word;
            text-size-adjust: none;
            white-space: pre-wrap;
        }
        h1 { font-size: 2rem; margin-bottom: 0.5rem; margin-top: 0; }
        h2 { font-size: 1.75rem; margin-top: 1rem; margin-bottom: 0.5rem; }
        h3 { font-size: 1.5rem; margin-top: 0.75rem; margin-bottom: 0.25rem; }
        h4 { font-size: 1.25rem; margin-top: 0.5rem; margin-bottom: 0.25rem; }
        p { margin: 0; font-size: 1rem; line-height: 1.6; }
        ul, ol { margin: 0.25rem 0; padding-left: 2rem; }
        li { margin: 0.125rem 0; }
        blockquote { 
            border-left: 4px solid #6366f1; 
            margin: 0.5rem 0; 
            padding: 0.25rem 1rem; 
            background: #f8fafc; 
            color: #475569;
        }
        pre { 
            background: #1e293b; 
            color: #e2e8f0; 
            padding: 0.75rem; 
            border-radius: 8px; 
            overflow-x: auto; 
            white-space: pre-wrap; /* Wrap long lines for PDF/Print */
            margin: 0.5rem 0;
        }
        code { 
            background: #f1f5f9; 
            padding: 0.125rem 0.375rem; 
            border-radius: 4px; 
            font-size: 0.875rem; 
        }
        pre code { background: transparent; padding: 0; }
        hr { border: none; border-top: 2px solid #e2e8f0; margin: 1rem 0; }
        table { border-collapse: collapse; width: 100%; margin: 0.5rem 0; }
        th, td { border: 1px solid #e2e8f0; padding: 0.5rem; text-align: left; }
        th { background: #f8fafc; font-weight: 600; }
        img { max-width: 100%; height: auto; border-radius: 8px; }
        a { color: #6366f1; text-decoration: none; }
    """ if embed_styles else ""
    
    # Generate continuous HTML content
    content_html = tiptap_to_html(doc.content, page_layout) if doc.content else ""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{safe_title}</title>
    <style>
{pdf_styles}
{base_styles}
    </style>
</head>
<body>
{content_html}
</body>
</html>"""


async def generate_pdf_from_captured_html(
    captured_html: str, title: str, page_layout: str = "a4"
) -> bytes:
    """Generate PDF from captured browser DOM HTML for pixel-perfect output.
    
    The captured HTML comes directly from the browser's ProseMirror editor,
    so it has the exact same text layout. We just wrap it with matching
    styles and send to Gotenberg.
    """
    page_format = PAGE_FORMATS.get(
        page_layout if page_layout in PAGE_FORMATS else "a4",
        PAGE_FORMATS["a4"]
    )
    margins = page_format['margins_in']
    safe_title = html_module.escape(title)
    
    # Styles that match the browser editor exactly:
    # - Arial font (same as notion-like-editor.scss)
    # - line-height 1.6 (same as paragraph-node.scss)
    # - margin 0 on p (same as notion-like-editor.scss margin-top: 0 !important)
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{safe_title}</title>
    <style>
        body {{
            margin: 0;
            padding: 0;
            font-family: Arial, sans-serif;
            font-size: 1rem;
            line-height: 1.6;
            color: #1e293b;
            /* Match browser globals from _variables.scss */
            text-rendering: optimizeLegibility;
            -webkit-font-smoothing: antialiased;
            overflow-wrap: break-word;
            text-size-adjust: none;
        }}
        /* Match ProseMirror styles from paragraph-node.scss */
        *, *::before, *::after {{ box-sizing: border-box; }}
        .ProseMirror, [data-node-type] {{ white-space: pre-wrap; }}
        p {{ margin: 0; font-size: 1rem; line-height: 1.6; }}
        h1 {{ font-size: 2rem; margin-bottom: 0.5rem; margin-top: 0; }}
        h2 {{ font-size: 1.75rem; margin-top: 1rem; margin-bottom: 0.5rem; }}
        h3 {{ font-size: 1.5rem; margin-top: 0.75rem; margin-bottom: 0.25rem; }}
        h4 {{ font-size: 1.25rem; margin-top: 0.5rem; margin-bottom: 0.25rem; }}
        ul, ol {{ margin: 0.25rem 0; padding-left: 2rem; }}
        li {{ margin: 0.125rem 0; }}
        blockquote {{
            border-left: 3px solid #e2e8f0;
            padding-left: 1rem;
            margin: 0.5rem 0;
            color: #64748b;
        }}
        pre {{
            background: #f1f5f9;
            padding: 1rem;
            border-radius: 0.375rem;
            overflow-x: auto;
            font-size: 0.875rem;
        }}
        img {{ max-width: 100%; display: block; margin: 10px 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 0.5rem 0; }}
        th, td {{ border: 1px solid #e2e8f0; padding: 0.5rem; text-align: left; }}
        th {{ background: #f8fafc; font-weight: 600; }}
    </style>
</head>
<body>
{captured_html}
</body>
</html>"""
    
    gotenberg_endpoint = f"{GOTENBERG_URL}/forms/chromium/convert/html"
    
    data = {
        "marginTop": str(margins['top']),
        "marginBottom": str(margins['bottom']),
        "marginLeft": str(margins['left']),
        "marginRight": str(margins['right']),
        "paperWidth": str(page_format['width_in']),
        "paperHeight": str(page_format['height_in']),
        "printBackground": "true",
        "scale": "1.0",
    }
    
    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(
            gotenberg_endpoint,
            data=data,
            files={"files": ("index.html", html_content.encode("utf-8"), "text/html")},
        )
        response.raise_for_status()
        return response.content


async def generate_document_pdf(doc: Any, page_layout: str = "document") -> bytes:
    """Generate PDF using Gotenberg's HTML-to-PDF conversion."""
    
    is_paginated = page_layout in PAGE_FORMATS
    page_format = PAGE_FORMATS.get(page_layout, PAGE_FORMATS["a4"])
    
    # Generate the full HTML document
    html_content = generate_document_html(doc, embed_styles=True, page_layout=page_layout, for_pdf=True)
    
    gotenberg_endpoint = f"{GOTENBERG_URL}/forms/chromium/convert/html"
    
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            # We configure Gotenberg with the exact paper dimensions and margins.
            # Gotenberg/Chromium will handle the content flow and pagination naturally.
            if is_paginated:
                margins = page_format['margins_in']
                data = {
                    "marginTop": str(margins['top']),
                    "marginBottom": str(margins['bottom']),
                    "marginLeft": str(margins['left']),
                    "marginRight": str(margins['right']),
                    "paperWidth": str(page_format['width_in']),
                    "paperHeight": str(page_format['height_in']),
                    "printBackground": "true",
                    "scale": "1.0",
                }
            else:
                # Default fallback for "full" or "document" layout
                data = {
                    "marginTop": "0.75",
                    "marginBottom": "0.75",
                    "marginLeft": "0.75",
                    "marginRight": "0.75",
                    "paperWidth": "8.27",
                    "paperHeight": "11.69",
                    "printBackground": "true",
                }
            
            response = await client.post(
                gotenberg_endpoint,
                files={"index.html": ("index.html", html_content.encode("utf-8"), "text/html")},
                data=data,
            )
            
            if response.status_code != 200:
                raise RuntimeError(f"Gotenberg PDF generation failed: {response.status_code} - {response.text}")
            
            return response.content
            
    except httpx.ConnectError as e:
        raise RuntimeError(f"Could not connect to Gotenberg at {GOTENBERG_URL}: {e}")
    except httpx.TimeoutException as e:
        raise RuntimeError(f"Gotenberg request timed out: {e}")


def generate_document_docx(doc: Any, page_layout: str = "document") -> bytes:
    """Generate Word document export with proper page sizing."""
    try:
        from docx import Document
        from docx.shared import Inches, Mm, Pt
        from docx.enum.text import WD_BREAK, WD_LINE_SPACING
    except ImportError:
        raise RuntimeError("python-docx is required for Word export. Install with: pip install python-docx")
    
    word_doc = Document()
    
    # Set default font to Arial 12pt to match browser editor
    style = word_doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(19.2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    # Disable "snap to grid" on Normal style — prevents Word from
    # rounding line heights to the document grid, which causes spacing drift.
    from docx.oxml.ns import qn
    pPr = style.element.get_or_add_pPr()
    snap = pPr.find(qn('w:snapToGrid'))
    if snap is None:
        snap = pPr.makeelement(qn('w:snapToGrid'), {qn('w:val'): '0'})
        pPr.append(snap)
    else:
        snap.set(qn('w:val'), '0')
    
    section = word_doc.sections[0]
    
    # Set page size based on layout (use PAGE_FORMATS for all known formats)
    if page_layout in PAGE_FORMATS:
        page_format = PAGE_FORMATS[page_layout]
        section.page_width = Inches(page_format["width_in"])
        section.page_height = Inches(page_format["height_in"])
        section.top_margin = Inches(page_format["margins_in"]["top"])
        section.bottom_margin = Inches(page_format["margins_in"]["bottom"])
        section.left_margin = Inches(page_format["margins_in"]["left"])
        section.right_margin = Inches(page_format["margins_in"]["right"])
    else:
        # Default to Letter for "document" view
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process TipTap content
    if doc.content:
        # We process the content as a single stream.
        # Word will handle the pagination automatically.
        nodes = doc.content.get("content", []) if isinstance(doc.content, dict) else []
        _add_nodes_to_docx(word_doc, nodes)
    
    buffer = io.BytesIO()
    word_doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_nodes_to_docx(word_doc: Any, nodes: List[Dict]) -> None:
    """Add a list of TipTap nodes to a Word document."""
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    
    ALIGN_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    
    def extract_text_from_node(node: Dict[str, Any]) -> str:
        if node.get("type") == "text":
            return node.get("text", "")
        texts = []
        for child in node.get("content", []):
            texts.append(extract_text_from_node(child))
        return "".join(texts)
    
    def add_runs_to_paragraph(p, content_nodes: List[Dict]) -> None:
        for node in content_nodes:
            if node.get("type") == "text":
                text = node.get("text", "")
                marks = node.get("marks", [])
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(12)  # 1rem = 16px = 12pt
                for mark in marks:
                    mark_type = mark.get("type", "")
                    if mark_type == "bold":
                        run.bold = True
                    elif mark_type == "italic":
                        run.italic = True
                    elif mark_type == "underline":
                        run.underline = True
                    elif mark_type == "strike":
                        run.font.strike = True
                    elif mark_type == "code":
                        run.font.name = 'Courier New'
            elif node.get("type") == "hardBreak":
                run = p.add_run("\n")
                run.font.name = 'Arial'
    
    def set_paragraph_spacing(p):
        """Match browser: CSS line-height 1.6 at 1rem (16px = 12pt).
        CSS line-height 1.6 = 1.6 × font-size = 1.6 × 12pt = 19.2pt.
        Word 'multiple 1.6' = 1.6 × single-spacing (~14.4pt) = ~23pt — WRONG.
        Must use exact line spacing instead."""
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(19.2)
    
    for node in nodes:
        node_type = node.get("type", "")
        
        if node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            text = extract_text_from_node(node)
            h = word_doc.add_heading(text, level)
            h.paragraph_format.space_before = Pt(12) if level == 1 else Pt(8)
            h.paragraph_format.space_after = Pt(4)
            # Set Arial on heading runs (override heading style default)
            for run in h.runs:
                run.font.name = 'Arial'
        
        elif node_type == "paragraph":
            p = word_doc.add_paragraph()
            add_runs_to_paragraph(p, node.get("content", []))
            set_paragraph_spacing(p)
            # Preserve text alignment from TipTap
            text_align = node.get("attrs", {}).get("textAlign") or "left"
            if text_align in ALIGN_MAP:
                p.paragraph_format.alignment = ALIGN_MAP[text_align]
        
        elif node_type == "bulletList":
            for item in node.get("content", []):
                if item.get("type") == "listItem":
                    for child in item.get("content", []):
                        text = extract_text_from_node(child)
                        p = word_doc.add_paragraph(text, style='List Bullet')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
        
        elif node_type == "orderedList":
            for item in node.get("content", []):
                if item.get("type") == "listItem":
                    for child in item.get("content", []):
                        text = extract_text_from_node(child)
                        p = word_doc.add_paragraph(text, style='List Number')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
        
        elif node_type == "codeBlock":
            text = extract_text_from_node(node)
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            set_paragraph_spacing(p)
        
        elif node_type == "blockquote":
            text = extract_text_from_node(node)
            p = word_doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            p.add_run(text).italic = True
            set_paragraph_spacing(p)
        
        elif node_type == "horizontalRule":
            p = word_doc.add_paragraph("─" * 50)
            set_paragraph_spacing(p)

        # Handle explicit user-inserted Page Breaks
        elif node_type == "pageBreak":
            p = word_doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
        
        elif node_type == "image":
            src = node.get("attrs", {}).get("src", "")
            p = word_doc.add_paragraph()
            p.add_run(f"[Image: {src}]").italic = True
            set_paragraph_spacing(p)
        
        elif node_type == "process-recording-node":
            session_id = node.get("attrs", {}).get("sessionId", "unknown")
            p = word_doc.add_paragraph()
            p.add_run(f"[Embedded Workflow Recording: {session_id}]").italic = True
            set_paragraph_spacing(p)
        
        elif node_type == "dataTable":
            table_name = node.get("attrs", {}).get("tableName", "")
            table_id = node.get("attrs", {}).get("tableId", "unknown")
            label = table_name or table_id
            p = word_doc.add_paragraph()
            p.add_run(f"[Embedded Data Table: {label}]").italic = True
            set_paragraph_spacing(p)
        
        elif node_type == "taskList":
            for item in node.get("content", []):
                if item.get("type") == "taskItem":
                    checked = item.get("attrs", {}).get("checked", False)
                    checkbox = "☑" if checked else "☐"
                    for child in item.get("content", []):
                        text = extract_text_from_node(child)
                        p = word_doc.add_paragraph(f"{checkbox} {text}", style='List Bullet')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)