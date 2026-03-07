"""
Utility functions for exporting workflows to various formats.
"""
import io
import base64
import os
import httpx
from typing import List, Dict, Any, Optional
from datetime import datetime

# Gotenberg configuration
GOTENBERG_URL = os.getenv("GOTENBERG_URL", "http://gotenberg:3000")


async def get_step_image_bytes(storage_path: str, file_path: str, storage_type: str = "local") -> Optional[bytes]:
    """Read step image bytes via the storage backend.

    Works with local filesystem, S3-compatible, GCS, and Azure — whatever
    the recording session was stored with.
    """
    from app.services.storage import get_storage_backend

    backend = get_storage_backend(storage_type)
    try:
        data = await backend.read_file(storage_path, file_path)
        if data:
            return data
    except Exception as e:
        print(f"[Export] Error reading image via backend ({storage_type}): {e}")
    return None


async def get_step_image_base64(storage_path: str, file_path: str, storage_type: str = "local") -> Optional[str]:
    """Get base64 encoded image for embedding in exports."""
    data = await get_step_image_bytes(storage_path, file_path, storage_type)
    if data:
        print(f"[Export] Read image, size: {len(data)} bytes")
        return base64.b64encode(data).decode("utf-8")
    return None


def generate_markdown(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    include_images: bool = False,
    image_base_url: Optional[str] = None,
) -> str:
    """Generate Markdown export of a workflow."""
    lines = []
    
    # Title
    title = workflow.get("name") or "Untitled Workflow"
    lines.append(f"# {title}")
    lines.append("")
    
    # Metadata
    created_at = workflow.get("created_at")
    if created_at:
        if isinstance(created_at, datetime):
            created_at = created_at.strftime("%Y-%m-%d %H:%M")
        lines.append(f"**Created:** {created_at}")
    
    total_steps = len([s for s in steps if s.get("step_type") in ("screenshot", "capture", "gif", "video", None)])
    lines.append(f"**Steps:** {total_steps}")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # Steps
    visible_index = 0
    for step in sorted(steps, key=lambda s: s.get("step_number", 0)):
        step_type = step.get("step_type") or "screenshot"
        step_number = step.get("step_number", 0)
        
        if step_type == "header":
            content = step.get("content") or step.get("description") or "Header"
            lines.append(f"## {content}")
            lines.append("")
        elif step_type == "tip":
            content = step.get("content") or step.get("description") or ""
            lines.append(f"> 💡 **Tip:** {content}")
            lines.append("")
        elif step_type == "alert":
            content = step.get("content") or step.get("description") or ""
            lines.append(f"> ⚠️ **Alert:** {content}")
            lines.append("")
        else:
            # Image-based step
            visible_index += 1
            description = step.get("description") or step.get("window_title") or f"Step {visible_index}"
            lines.append(f"### Step {visible_index}: {description}")
            lines.append("")
            
            # Add image if available
            if include_images and step_number in files:
                if image_base_url:
                    img_url = f"{image_base_url}/session/{workflow['id']}/image/{step_number}"
                    lines.append(f"![Step {visible_index}]({img_url})")
                else:
                    lines.append(f"*[Image for step {visible_index}]*")
                lines.append("")
            
            # Additional details
            if step.get("text_typed"):
                lines.append(f"**Text entered:** `{step['text_typed']}`")
                lines.append("")
            if step.get("key_pressed"):
                lines.append(f"**Key pressed:** `{step['key_pressed']}`")
                lines.append("")
    
    return "\n".join(lines)


async def generate_html(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    storage_path: str,
    storage_type: str = "local",
    embed_images: bool = True,
    image_base_url: Optional[str] = None,
    for_pdf: bool = False,  # NEW: flag for PDF-optimized HTML
) -> str:
    """Generate HTML export of a workflow."""
    title = workflow.get("name") or "Untitled Workflow"
    created_at = workflow.get("created_at")
    if isinstance(created_at, datetime):
        created_at = created_at.strftime("%Y-%m-%d %H:%M")
    
    total_steps = len([s for s in steps if s.get("step_type") in ("screenshot", "capture", "gif", "video", None)])
    
    # PDF-specific styles for better page breaks
    pdf_styles = """
        @page { 
            size: A4; 
            margin: 1.5cm; 
        }
        @media print {
            .step { page-break-inside: avoid; }
            .section-header { page-break-after: avoid; }
        }
    """ if for_pdf else ""
    
    html_parts = [
        "<!DOCTYPE html>",
        "<html lang='en'>",
        "<head>",
        "  <meta charset='UTF-8'>",
        "  <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        f"  <title>{title}</title>",
        "  <style>",
        f"    {pdf_styles}",
        "    body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 900px; margin: 0 auto; padding: 2rem; background: #f8fafc; }",
        "    h1 { color: #1e293b; border-bottom: 2px solid #e2e8f0; padding-bottom: 1rem; }",
        "    .meta { color: #64748b; margin-bottom: 2rem; }",
        "    .step { background: white; border-radius: 12px; padding: 1.5rem; margin-bottom: 1.5rem; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }",
        "    .step-header { display: flex; align-items: center; gap: 0.75rem; margin-bottom: 1rem; }",
        "    .step-number { background: #6366f1; color: white; width: 2rem; height: 2rem; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: 600; font-size: 0.875rem; }",
        "    .step-title { font-size: 1.125rem; font-weight: 600; color: #1e293b; }",
        "    .step img { max-width: 100%; border-radius: 8px; border: 1px solid #e2e8f0; }",
        "    .tip { background: #ecfdf5; border-left: 4px solid #10b981; padding: 1rem; border-radius: 0 8px 8px 0; margin-bottom: 1.5rem; }",
        "    .alert { background: #fef3c7; border-left: 4px solid #f59e0b; padding: 1rem; border-radius: 0 8px 8px 0; margin-bottom: 1.5rem; }",
        "    .section-header { font-size: 1.5rem; font-weight: 700; color: #1e293b; margin: 2rem 0 1rem; }",
        "    .detail { color: #64748b; font-size: 0.875rem; margin-top: 0.5rem; }",
        "    code { background: #f1f5f9; padding: 0.125rem 0.375rem; border-radius: 4px; font-size: 0.875rem; }",
        "    .no-image { background: #f1f5f9; padding: 2rem; text-align: center; color: #64748b; border-radius: 8px; }",
        "  </style>",
        "</head>",
        "<body>",
        f"  <h1>{title}</h1>",
        f"  <div class='meta'>Created: {created_at or 'Unknown'} • {total_steps} steps</div>",
    ]
    
    # Debug: log available files
    print(f"[PDF Export] Storage path: {storage_path}")
    print(f"[PDF Export] Files dict: {files}")
    
    visible_index = 0
    for step in sorted(steps, key=lambda s: s.get("step_number", 0)):
        step_type = step.get("step_type") or "screenshot"
        step_number = step.get("step_number", 0)
        
        if step_type == "header":
            content = step.get("content") or step.get("description") or "Header"
            html_parts.append(f"  <h2 class='section-header'>{content}</h2>")
        elif step_type == "tip":
            content = step.get("content") or step.get("description") or ""
            html_parts.append(f"  <div class='tip'>💡 <strong>Tip:</strong> {content}</div>")
        elif step_type == "alert":
            content = step.get("content") or step.get("description") or ""
            html_parts.append(f"  <div class='alert'>⚠️ <strong>Alert:</strong> {content}</div>")
        else:
            visible_index += 1
            description = step.get("description") or step.get("window_title") or f"Step {visible_index}"
            
            html_parts.append("  <div class='step'>")
            html_parts.append("    <div class='step-header'>")
            html_parts.append(f"      <div class='step-number'>{visible_index}</div>")
            html_parts.append(f"      <div class='step-title'>{description}</div>")
            html_parts.append("    </div>")
            
            # Add image
            if step_number in files:
                file_path = files[step_number]
                print(f"[PDF Export] Step {step_number} has file: {file_path}")
                
                if embed_images:
                    img_b64 = await get_step_image_base64(storage_path, file_path, storage_type)
                    if img_b64:
                        html_parts.append(f"    <img src='data:image/png;base64,{img_b64}' alt='Step {visible_index}'>")
                    else:
                        html_parts.append(f"    <div class='no-image'>[Image not available: {file_path}]</div>")
                elif image_base_url:
                    img_url = f"{image_base_url}/session/{workflow['id']}/image/{step_number}"
                    html_parts.append(f"    <img src='{img_url}' alt='Step {visible_index}'>")
            else:
                print(f"[PDF Export] Step {step_number} has no file in files dict")
            
            # Additional details
            if step.get("text_typed"):
                html_parts.append(f"    <div class='detail'>Text entered: <code>{step['text_typed']}</code></div>")
            if step.get("key_pressed"):
                html_parts.append(f"    <div class='detail'>Key pressed: <code>{step['key_pressed']}</code></div>")
            
            html_parts.append("  </div>")
    
    html_parts.extend([
        "</body>",
        "</html>",
    ])
    
    return "\n".join(html_parts)


async def generate_pdf_gotenberg(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    storage_path: str,
    storage_type: str = "local",
) -> bytes:
    """Generate PDF using Gotenberg's HTML-to-PDF conversion."""
    
    # Generate HTML optimized for PDF
    html_content = await generate_html(
        workflow,
        steps,
        files,
        storage_path,
        storage_type,
        embed_images=True,  # Always embed for PDF
        for_pdf=True,
    )
    
    # Call Gotenberg API
    gotenberg_endpoint = f"{GOTENBERG_URL}/forms/chromium/convert/html"
    
    async with httpx.AsyncClient(timeout=120.0) as client:
        # Gotenberg expects multipart form data
        files_data = {
            "index.html": ("index.html", html_content.encode("utf-8"), "text/html"),
        }
        
        # PDF options
        data = {
            "marginTop": "0.5",
            "marginBottom": "0.5",
            "marginLeft": "0.5",
            "marginRight": "0.5",
            "paperWidth": "8.27",  # A4 width in inches
            "paperHeight": "11.69",  # A4 height in inches
            "printBackground": "true",
        }
        
        response = await client.post(
            gotenberg_endpoint,
            files=files_data,
            data=data,
        )
        
        if response.status_code != 200:
            raise RuntimeError(f"Gotenberg PDF generation failed: {response.status_code} - {response.text}")
        
        return response.content


def generate_pdf(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    storage_path: str,
    storage_type: str = "local",
) -> bytes:
    """
    Synchronous wrapper for PDF generation.
    Falls back to reportlab if Gotenberg is unavailable.
    """
    import asyncio
    
    try:
        # Try Gotenberg first
        return asyncio.get_event_loop().run_until_complete(
            generate_pdf_gotenberg(workflow, steps, files, storage_path, storage_type)
        )
    except Exception as e:
        # Fall back to reportlab if available
        try:
            return _generate_pdf_reportlab(workflow, steps, files, storage_path, storage_type)
        except ImportError:
            raise RuntimeError(f"PDF generation failed. Gotenberg error: {e}. Install reportlab as fallback: pip install reportlab")


def _generate_pdf_reportlab(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    storage_path: str,
    storage_type: str = "local",
) -> bytes:
    """Fallback PDF generation using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        raise ImportError("reportlab is required for fallback PDF export")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=20,
        alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=10,
    )
    step_title_style = ParagraphStyle(
        'StepTitle',
        parent=styles['Heading3'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#1e293b'),
    )
    normal_style = styles['Normal']
    tip_style = ParagraphStyle(
        'Tip',
        parent=normal_style,
        backColor=colors.HexColor('#ecfdf5'),
        borderPadding=10,
        spaceBefore=10,
        spaceAfter=10,
    )
    alert_style = ParagraphStyle(
        'Alert',
        parent=normal_style,
        backColor=colors.HexColor('#fef3c7'),
        borderPadding=10,
        spaceBefore=10,
        spaceAfter=10,
    )
    
    story = []
    title = workflow.get("name") or "Untitled Workflow"
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 20))
    
    # Simplified version - add steps
    visible_index = 0
    for step in sorted(steps, key=lambda s: s.get("step_number", 0)):
        step_type = step.get("step_type") or "screenshot"
        if step_type in ("screenshot", "capture", "gif", "video"):
            visible_index += 1
            description = step.get("description") or f"Step {visible_index}"
            story.append(Paragraph(f"Step {visible_index}: {description}", styles['Heading2']))
            story.append(Spacer(1, 10))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_confluence_storage(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    image_base_url: Optional[str] = None,
) -> str:
    """Generate Confluence Storage Format (XML-based) export of a workflow."""
    title = workflow.get("name") or "Untitled Workflow"
    created_at = workflow.get("created_at")
    if isinstance(created_at, datetime):
        created_at = created_at.strftime("%Y-%m-%d %H:%M")

    total_steps = len([s for s in steps if s.get("step_type") in ("screenshot", "capture", "gif", "video", None)])

    parts = []
    parts.append(f"<h1>{title}</h1>")
    parts.append(f'<ac:structured-macro ac:name="panel"><ac:rich-text-body>')
    parts.append(f"<p><strong>Created:</strong> {created_at or 'Unknown'} &bull; <strong>Steps:</strong> {total_steps}</p>")
    parts.append("</ac:rich-text-body></ac:structured-macro>")
    parts.append("<hr/>")

    visible_index = 0
    for step in sorted(steps, key=lambda s: s.get("step_number", 0)):
        step_type = step.get("step_type") or "screenshot"
        step_number = step.get("step_number", 0)

        if step_type == "header":
            content = step.get("content") or step.get("description") or "Header"
            parts.append(f"<h2>{content}</h2>")
        elif step_type == "tip":
            content = step.get("content") or step.get("description") or ""
            parts.append(f'<ac:structured-macro ac:name="tip"><ac:rich-text-body><p>{content}</p></ac:rich-text-body></ac:structured-macro>')
        elif step_type == "alert":
            content = step.get("content") or step.get("description") or ""
            parts.append(f'<ac:structured-macro ac:name="warning"><ac:rich-text-body><p>{content}</p></ac:rich-text-body></ac:structured-macro>')
        else:
            visible_index += 1
            description = step.get("description") or step.get("window_title") or f"Step {visible_index}"
            parts.append(f'<ac:structured-macro ac:name="expand"><ac:parameter ac:name="title">Step {visible_index}: {description}</ac:parameter><ac:rich-text-body>')

            if step_number in files and image_base_url:
                img_url = f"{image_base_url}/session/{workflow['id']}/image/{step_number}"
                parts.append(f'<p><ac:image><ri:url ri:value="{img_url}"/></ac:image></p>')
            elif step_number in files:
                filename = os.path.basename(files[step_number])
                parts.append(f'<p><ac:image><ri:attachment ri:filename="{filename}"/></ac:image></p>')

            if step.get("text_typed"):
                parts.append(f"<p><strong>Text entered:</strong> <code>{step['text_typed']}</code></p>")
            if step.get("key_pressed"):
                parts.append(f"<p><strong>Key pressed:</strong> <code>{step['key_pressed']}</code></p>")

            parts.append("</ac:rich-text-body></ac:structured-macro>")

    return "\n".join(parts)


def generate_notion_markdown(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    image_base_url: Optional[str] = None,
) -> str:
    """Generate Notion-compatible Markdown export of a workflow."""
    title = workflow.get("name") or "Untitled Workflow"
    created_at = workflow.get("created_at")
    if isinstance(created_at, datetime):
        created_at = created_at.strftime("%Y-%m-%d %H:%M")

    total_steps = len([s for s in steps if s.get("step_type") in ("screenshot", "capture", "gif", "video", None)])

    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"> 📋 **Created:** {created_at or 'Unknown'} • **Steps:** {total_steps}")
    lines.append("")
    lines.append("---")
    lines.append("")

    visible_index = 0
    for step in sorted(steps, key=lambda s: s.get("step_number", 0)):
        step_type = step.get("step_type") or "screenshot"
        step_number = step.get("step_number", 0)

        if step_type == "header":
            content = step.get("content") or step.get("description") or "Header"
            lines.append(f"## {content}")
            lines.append("")
        elif step_type == "tip":
            content = step.get("content") or step.get("description") or ""
            lines.append(f"> 💡 **Tip:** {content}")
            lines.append("")
        elif step_type == "alert":
            content = step.get("content") or step.get("description") or ""
            lines.append(f"> ⚠️ **Alert:** {content}")
            lines.append("")
        else:
            visible_index += 1
            description = step.get("description") or step.get("window_title") or f"Step {visible_index}"

            lines.append(f"<details><summary><strong>Step {visible_index}:</strong> {description}</summary>")
            lines.append("")

            if step_number in files and image_base_url:
                img_url = f"{image_base_url}/session/{workflow['id']}/image/{step_number}"
                lines.append(f"![Step {visible_index}]({img_url})")
                lines.append("")

            if step.get("text_typed"):
                lines.append(f"**Text entered:** `{step['text_typed']}`")
                lines.append("")
            if step.get("key_pressed"):
                lines.append(f"**Key pressed:** `{step['key_pressed']}`")
                lines.append("")

            lines.append("</details>")
            lines.append("")
            lines.append("---")
            lines.append("")

    return "\n".join(lines)


async def generate_docx(
    workflow: Dict[str, Any],
    steps: List[Dict[str, Any]],
    files: Dict[int, str],
    storage_path: str,
    storage_type: str = "local",
) -> bytes:
    """Generate Word document export of a workflow."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is required for Word export. Install with: pip install python-docx")
    
    doc = Document()
    
    # Title
    title = workflow.get("name") or "Untitled Workflow"
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    created_at = workflow.get("created_at")
    if isinstance(created_at, datetime):
        created_at = created_at.strftime("%Y-%m-%d %H:%M")
    
    total_steps = len([s for s in steps if s.get("step_type") in ("screenshot", "capture", "gif", "video", None)])
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Created: {created_at or 'Unknown'} • {total_steps} steps").italic = True
    
    doc.add_paragraph()
    
    # Debug logging
    print(f"[DOCX Export] Storage path: {storage_path}")
    print(f"[DOCX Export] Files dict: {files}")
    
    # Steps
    visible_index = 0
    for step in sorted(steps, key=lambda s: s.get("step_number", 0)):
        step_type = step.get("step_type") or "screenshot"
        step_number = step.get("step_number", 0)
        
        if step_type == "header":
            content = step.get("content") or step.get("description") or "Header"
            doc.add_heading(content, 1)
        elif step_type == "tip":
            content = step.get("content") or step.get("description") or ""
            p = doc.add_paragraph()
            p.add_run("💡 Tip: ").bold = True
            p.add_run(content)
        elif step_type == "alert":
            content = step.get("content") or step.get("description") or ""
            p = doc.add_paragraph()
            p.add_run("⚠️ Alert: ").bold = True
            p.add_run(content)
        else:
            visible_index += 1
            description = step.get("description") or step.get("window_title") or f"Step {visible_index}"
            
            doc.add_heading(f"Step {visible_index}: {description}", 2)
            
            # Add image
            if step_number in files:
                file_path = files[step_number]
                print(f"[DOCX Export] Step {step_number} has file: {file_path}")
                
                # Read image via storage backend (works with S3, GCS, Azure, local)
                img_data = await get_step_image_bytes(storage_path, file_path, storage_type)
                
                if img_data:
                    try:
                        img_stream = io.BytesIO(img_data)
                        doc.add_picture(img_stream, width=Inches(6))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"[DOCX Export] Added image for step {step_number}, size: {len(img_data)} bytes")
                    except Exception as e:
                        print(f"[DOCX Export] Error adding image: {e}")
                        doc.add_paragraph(f"[Image could not be loaded: {e}]")
                else:
                    print(f"[DOCX Export] Image not found for step {step_number}")
                    doc.add_paragraph(f"[Image not available: {file_path}]")
            else:
                print(f"[DOCX Export] Step {step_number} has no file in files dict")
            
            # Additional details
            if step.get("text_typed"):
                p = doc.add_paragraph()
                p.add_run("Text entered: ").bold = True
                p.add_run(step['text_typed'])
            if step.get("key_pressed"):
                p = doc.add_paragraph()
                p.add_run("Key pressed: ").bold = True
                p.add_run(step['key_pressed'])
            
            doc.add_paragraph()
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
